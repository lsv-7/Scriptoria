import os
import io
import datetime
import json
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def safe_str(val, default="N/A"):
    if val is None:
        return default
    if isinstance(val, dict):
        for k in ["description", "details", "text", "info"]:
            if k in val and isinstance(val[k], str):
                return val[k]
        return json.dumps(val, indent=2)
    if isinstance(val, list):
        return ", ".join([str(item) for item in val])
    return str(val)

# --- CUSTOM REPORTLAB CANVAS FOR DYNAMIC PAGE NUMBERING ---
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        
        # Suppress headers/footers on page 1 (Cover Page)
        if self._pageNumber > 1:
            # Header
            self.setFont("Helvetica-Bold", 8)
            self.setFillColor(colors.HexColor("#D4AF37")) # Gold accent
            self.drawString(54, 755, "CINEFORGE AI")
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#4a4a5a"))
            self.drawString(125, 755, "|   PRE-PRODUCTION BLUEPRINT")
            
            # Header line
            self.setStrokeColor(colors.HexColor("#dcdce6"))
            self.setLineWidth(0.5)
            self.line(54, 745, 558, 745)
            
            # Footer line
            self.line(54, 55, 558, 55)
            
            # Footer text
            self.setFont("Helvetica", 8)
            self.drawString(54, 40, "Confidential - Generated via CineForge AI Studio")
            page_text = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(558, 40, page_text)
            
        self.restoreState()


class ExportService:
    def generate_pdf(self, project, data):
        """
        Generates a professionally-styled ReportLab PDF document in memory.
        Returns bytes.
        """
        buffer = io.BytesIO()
        
        # Setup document geometry (letter size, 0.75 in / 54pt margins)
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=72,
            bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        # Custom color definitions (matching landing page styling)
        primary_color = colors.HexColor("#0A0A0A") # Obsidian Black
        accent_color = colors.HexColor("#D4AF37")  # Gold
        secondary_color = colors.HexColor("#F5C542") # Bright Gold
        body_color = colors.HexColor("#2e2e3a")
        
        # Custom typography styles
        styles.add(ParagraphStyle(
            name='CoverTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=32,
            leading=38,
            textColor=primary_color,
            alignment=1, # Centered
            spaceAfter=15
        ))
        
        styles.add(ParagraphStyle(
            name='CoverSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=18,
            textColor=accent_color,
            alignment=1,
            spaceAfter=40
        ))
        
        styles.add(ParagraphStyle(
            name='CoverMeta',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=body_color,
            alignment=1
        ))

        styles.add(ParagraphStyle(
            name='SectionHeading',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=accent_color,
            spaceBefore=18,
            spaceAfter=10,
            keepWithNext=True
        ))

        styles.add(ParagraphStyle(
            name='SubsectionHeading',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=primary_color,
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True
        ))

        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=body_color,
            spaceAfter=8
        ))

        styles.add(ParagraphStyle(
            name='ScreenplayText',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=10,
            leading=14,
            textColor=colors.black,
            leftIndent=36,
            spaceAfter=6
        ))

        styles.add(ParagraphStyle(
            name='ScreenplayDialogue',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=10,
            leading=14,
            textColor=colors.black,
            leftIndent=72,
            rightMargin=72,
            spaceAfter=6
        ))

        styles.add(ParagraphStyle(
            name='ScreenplayCharacter',
            parent=styles['Normal'],
            fontName='Courier-Bold',
            fontSize=10,
            leading=14,
            textColor=colors.black,
            leftIndent=144,
            spaceAfter=2,
            keepWithNext=True
        ))

        story = []
        
        # --- COVER PAGE ---
        story.append(Spacer(1, 100))
        story.append(Paragraph("CINEFORGE AI", styles['CoverSubtitle']))
        story.append(Paragraph(project.get("project_name", "Untitled Project").upper(), styles['CoverTitle']))
        story.append(Paragraph("PRE-PRODUCTION STUDIO BLUEPRINT", styles['CoverSubtitle']))
        story.append(Spacer(1, 100))
        
        # Cover metadata box
        meta_html = f"""
        <b>Genre:</b> {project.get('genre', 'N/A')}<br/>
        <b>Target Audience:</b> {project.get('target_audience', 'N/A')}<br/>
        <b>Date Created:</b> {datetime.datetime.now().strftime('%B %d, %Y')}<br/>
        """
        story.append(Paragraph(meta_html, styles['CoverMeta']))
        story.append(PageBreak())
        
        # --- 1. STORY ANALYSIS ---
        analysis = data.get("story_analysis", {})
        story.append(Paragraph("1. STORY ANALYSIS", styles['SectionHeading']))
        story.append(Paragraph(f"<b>Logline:</b> {safe_str(analysis.get('logline', 'N/A'))}", styles['CustomBody']))
        story.append(Paragraph(f"<b>Tagline:</b> <i>\"{safe_str(analysis.get('tagline', 'N/A'))}\"</i>", styles['CustomBody']))
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("Synopsis", styles['SubsectionHeading']))
        synopsis = safe_str(analysis.get('synopsis', 'N/A'))
        for para in synopsis.split('\n\n'):
            story.append(Paragraph(para, styles['CustomBody']))
            
        story.append(Paragraph("Genre Analysis", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(analysis.get('genre_analysis', 'N/A')), styles['CustomBody']))
        
        story.append(Paragraph("Theme & Core Message", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(analysis.get('theme', 'N/A')), styles['CustomBody']))
        
        story.append(Paragraph("Audience Insights", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(analysis.get('audience_insights', 'N/A')), styles['CustomBody']))
        story.append(PageBreak())

        # --- 2. NARRATIVE STRUCTURE ---
        struct = data.get("narrative_structure", {})
        story.append(Paragraph("2. NARRATIVE STRUCTURE (3-ACT BREAKDOWN)", styles['SectionHeading']))
        for act_key in ["act_1", "act_2", "act_3"]:
            act = struct.get(act_key, {})
            if act:
                story.append(Paragraph(safe_str(act.get("title", act_key.replace("_", " ").title())), styles['SubsectionHeading']))
                story.append(Paragraph(f"<b>Description:</b> {safe_str(act.get('description', 'N/A'))}", styles['CustomBody']))
                story.append(Paragraph(f"<b>Conflict:</b> {safe_str(act.get('conflict', 'N/A'))}", styles['CustomBody']))
                story.append(Paragraph(f"<b>Rising Action / Progression:</b> {safe_str(act.get('rising_action', 'N/A'))}", styles['CustomBody']))
                if "resolution" in act:
                    story.append(Paragraph(f"<b>Resolution:</b> {safe_str(act.get('resolution', 'N/A'))}", styles['CustomBody']))
                story.append(Spacer(1, 10))
        story.append(PageBreak())

        # --- 3. CHARACTERS ---
        chars_data = data.get("characters", {}).get("characters", [])
        story.append(Paragraph("3. CHARACTER PROFILES", styles['SectionHeading']))
        for char in chars_data:
            story.append(Paragraph(f"<b>{safe_str(char.get('name', 'Unnamed'))}</b> (Age: {safe_str(char.get('age', 'N/A'))})", styles['SubsectionHeading']))
            story.append(Paragraph(f"<b>Backstory:</b> {safe_str(char.get('backstory', 'N/A'))}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Personality:</b> {safe_str(char.get('personality', 'N/A'))}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Core Goals:</b> {safe_str(char.get('goals', 'N/A'))}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Strengths:</b> {safe_str(char.get('strengths', 'N/A'))} | <b>Weaknesses:</b> {safe_str(char.get('weaknesses', 'N/A'))}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Character Arc:</b> {safe_str(char.get('arc', 'N/A'))}", styles['CustomBody']))
            story.append(Spacer(1, 12))
        story.append(PageBreak())

        # --- 4. SCENE BREAKDOWN ---
        scenes = data.get("scene_breakdown", {}).get("scenes", [])
        story.append(Paragraph("4. SCENE BREAKDOWN", styles['SectionHeading']))
        if scenes:
            table_data = [["Scene", "Location Heading", "Characters", "Objective", "Est. Dur"]]
            for s in scenes:
                chars_val = s.get("characters", "")
                if isinstance(chars_val, list):
                    chars_str = ", ".join([str(c) for c in chars_val])
                else:
                    chars_str = str(chars_val or "")
                table_data.append([
                    str(s.get("scene_number", "")),
                    s.get("location", ""),
                    chars_str,
                    s.get("objective", ""),
                    s.get("duration", "")
                ])
            # Draw table
            t = Table(table_data, colWidths=[40, 140, 110, 170, 40])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), accent_color),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                ('TOPPADDING', (0,0), (-1,-1), 6),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cccccc")),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#f8f8f8")])
            ]))
            story.append(t)
        else:
            story.append(Paragraph("No scenes broken down.", styles['CustomBody']))
        story.append(PageBreak())

        # --- 5. STORYBOARD PLAN ---
        storyboards = data.get("storyboard", {}).get("storyboards", [])
        story.append(Paragraph("5. STORYBOARD PLAN", styles['SectionHeading']))
        for sb in storyboards:
            story.append(Paragraph(f"Scene {sb.get('scene_number', 'N/A')} - Frame details", styles['SubsectionHeading']))
            story.append(Paragraph(f"<b>Camera Angle:</b> {sb.get('camera_angle', 'N/A')}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Lighting:</b> {sb.get('lighting', 'N/A')}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Mood/Tone:</b> {sb.get('mood', 'N/A')}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Art Prompt:</b> <i>\"{sb.get('prompt', 'N/A')}\"</i>", styles['CustomBody']))
            story.append(Spacer(1, 10))
        story.append(PageBreak())

        # --- 6. SOUND DESIGN PLANNER ---
        sound = data.get("sound_design", {})
        story.append(Paragraph("6. SOUND DESIGN PLAN", styles['SectionHeading']))
        story.append(Paragraph("Background Music / Soundtrack", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(sound.get("background_music")), styles['CustomBody']))
        story.append(Paragraph("Ambience & Atmosphere", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(sound.get("ambience")), styles['CustomBody']))
        story.append(Paragraph("Foley & Sound Effects", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(sound.get("foley_effects")), styles['CustomBody']))
        story.append(Paragraph("Dialogue Treatment", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(sound.get("dialogue_treatment")), styles['CustomBody']))
        story.append(Paragraph("Sound Cue Production Notes", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(sound.get("scene_sound_notes")), styles['CustomBody']))
        story.append(PageBreak())

        # --- 7. PRODUCTION PLANNING ---
        prod = data.get("production_plan", {})
        story.append(Paragraph("7. PRODUCTION PLANNING SHEET", styles['SectionHeading']))
        story.append(Paragraph("Shooting Locations", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(prod.get("shooting_locations")), styles['CustomBody']))
        story.append(Paragraph("Required Props", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(prod.get("required_props")), styles['CustomBody']))
        story.append(Paragraph("Production Equipment Suggested", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(prod.get("equipment")), styles['CustomBody']))
        story.append(Paragraph("Crew Suggestions", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(prod.get("crew_suggestions")), styles['CustomBody']))
        story.append(Paragraph("Estimated Shoot Days", styles['SubsectionHeading']))
        story.append(Paragraph(safe_str(prod.get("estimated_shoot_days")), styles['CustomBody']))
        story.append(PageBreak())

        # --- 8. BUDGET PLANNING ---
        budget = data.get("budget_plan", {})
        story.append(Paragraph("8. BUDGET PLANNING", styles['SectionHeading']))
        if budget:
            pre = budget.get("pre_production", {}) or {}
            prod_b = budget.get("production", {}) or {}
            post = budget.get("post_production", {}) or {}
            
            story.append(Paragraph("Pre-Production Budget", styles['SubsectionHeading']))
            story.append(Paragraph(f"<b>Estimated Cost:</b> {safe_str(pre.get('cost'))}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Details:</b> {safe_str(pre.get('details'))}", styles['CustomBody']))
            story.append(Spacer(1, 6))
            
            story.append(Paragraph("Production Budget", styles['SubsectionHeading']))
            story.append(Paragraph(f"<b>Estimated Cost:</b> {safe_str(prod_b.get('cost'))}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Details:</b> {safe_str(prod_b.get('details'))}", styles['CustomBody']))
            story.append(Spacer(1, 6))
            
            story.append(Paragraph("Post-Production Budget", styles['SubsectionHeading']))
            story.append(Paragraph(f"<b>Estimated Cost:</b> {safe_str(post.get('cost'))}", styles['CustomBody']))
            story.append(Paragraph(f"<b>Details:</b> {safe_str(post.get('details'))}", styles['CustomBody']))
            story.append(Spacer(1, 6))
            
            story.append(Paragraph(f"<b>Estimated Total Budget:</b> {safe_str(budget.get('total_budget'))}", styles['SubsectionHeading']))
            story.append(Paragraph(f"<b>Cost Saving Tips:</b> {safe_str(budget.get('cost_saving_tips'))}", styles['CustomBody']))
        else:
            story.append(Paragraph("No budget details generated.", styles['CustomBody']))
        story.append(PageBreak())

        # --- 9. SCREENPLAY SCRIPT EXCERPT ---
        screenplay = safe_str(data.get("screenplay", {}).get("screenplay_text", ""))
        story.append(Paragraph("9. SCREENPLAY SCRIPT EXCERPT", styles['SectionHeading']))
        if screenplay:
            lines = screenplay.split('\n')
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    story.append(Spacer(1, 6))
                    continue
                
                # Check formatting tags to map indentation (Screenplay style)
                # Headings INT. / EXT.
                if stripped.startswith("INT.") or stripped.startswith("EXT."):
                    story.append(Paragraph(stripped, styles['ScreenplayText']))
                # Dialogues (if centered or capital inside screenplay lines)
                # Screenplay lines have specific shapes
                # A simple script check: dialogue heading has uppercase with no punctuation
                elif stripped.isupper() and len(stripped) < 30 and not (stripped.startswith("INT.") or stripped.startswith("EXT.") or stripped.startswith("FADE ")):
                    story.append(Paragraph(stripped, styles['ScreenplayCharacter']))
                elif stripped.startswith("("):
                    # Parenthetical
                    p_style = ParagraphStyle('Parenthetical', parent=styles['ScreenplayDialogue'], fontName='Courier-Oblique', leftIndent=108)
                    story.append(Paragraph(stripped, p_style))
                elif len(line) - len(line.lstrip()) > 10:
                    # indented, likely dialogue text
                    story.append(Paragraph(stripped, styles['ScreenplayDialogue']))
                else:
                    # Action line
                    act_style = ParagraphStyle('ActionLine', parent=styles['ScreenplayText'], leftIndent=0)
                    story.append(Paragraph(stripped, act_style))
        else:
            story.append(Paragraph("No screenplay generated yet.", styles['CustomBody']))

        # Build Document
        doc.build(story, canvasmaker=NumberedCanvas)
        return buffer.getvalue()

    def generate_docx(self, project, data):
        """
        Generates a Word Document (DOCX) using python-docx.
        Returns bytes.
        """
        doc = Document()
        
        # Styles config
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Arial'
        style_normal.font.size = Pt(11)
        
        # Cover/Header
        title = doc.add_paragraph()
        run_title = title.add_run(project.get("project_name", "Untitled Project").upper() + "\n")
        run_title.font.size = Pt(28)
        run_title.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        run_sub = subtitle.add_run("PRE-PRODUCTION STUDIO BLUEPRINT\nGenerated by CineForge AI")
        run_sub.font.size = Pt(14)
        run_sub.font.italic = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Genre: {project.get('genre', 'N/A')}")
        doc.add_paragraph(f"Target Audience: {project.get('target_audience', 'N/A')}")
        doc.add_paragraph(f"Generated On: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()
        
        # 1. Story Analysis
        doc.add_heading("1. Story Analysis", level=1)
        analysis = data.get("story_analysis", {})
        doc.add_paragraph(f"Logline: {safe_str(analysis.get('logline', 'N/A'))}")
        doc.add_paragraph(f"Tagline: {safe_str(analysis.get('tagline', 'N/A'))}")
        doc.add_heading("Synopsis", level=2)
        doc.add_paragraph(safe_str(analysis.get('synopsis', 'N/A')))
        doc.add_heading("Genre Analysis", level=2)
        doc.add_paragraph(safe_str(analysis.get('genre_analysis', 'N/A')))
        doc.add_heading("Theme", level=2)
        doc.add_paragraph(safe_str(analysis.get('theme', 'N/A')))
        doc.add_heading("Audience Insights", level=2)
        doc.add_paragraph(safe_str(analysis.get('audience_insights', 'N/A')))
        
        # 2. Narrative Structure
        doc.add_heading("2. Narrative Structure", level=1)
        struct = data.get("narrative_structure", {})
        for act_key in ["act_1", "act_2", "act_3"]:
            act = struct.get(act_key, {})
            if act:
                doc.add_heading(safe_str(act.get("title", act_key.replace("_", " ").title())), level=2)
                doc.add_paragraph(f"Description: {safe_str(act.get('description', 'N/A'))}")
                doc.add_paragraph(f"Conflict: {safe_str(act.get('conflict', 'N/A'))}")
                doc.add_paragraph(f"Rising Action: {safe_str(act.get('rising_action', 'N/A'))}")
                if "resolution" in act:
                    doc.add_paragraph(f"Resolution: {safe_str(act.get('resolution', 'N/A'))}")

        # 3. Characters
        doc.add_heading("3. Character Profiles", level=1)
        chars = data.get("characters", {}).get("characters", [])
        for char in chars:
            doc.add_heading(safe_str(char.get("name", "Unnamed")), level=2)
            doc.add_paragraph(f"Age: {safe_str(char.get('age', 'N/A'))}")
            doc.add_paragraph(f"Backstory: {safe_str(char.get('backstory', 'N/A'))}")
            doc.add_paragraph(f"Personality: {safe_str(char.get('personality', 'N/A'))}")
            doc.add_paragraph(f"Goals: {safe_str(char.get('goals', 'N/A'))}")
            doc.add_paragraph(f"Strengths: {safe_str(char.get('strengths', 'N/A'))}")
            doc.add_paragraph(f"Weaknesses: {safe_str(char.get('weaknesses', 'N/A'))}")
            doc.add_paragraph(f"Character Arc: {safe_str(char.get('arc', 'N/A'))}")

        # 4. Scene Breakdown
        doc.add_heading("4. Scene Breakdown", level=1)
        scenes = data.get("scene_breakdown", {}).get("scenes", [])
        if scenes:
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Scene'
            hdr_cells[1].text = 'Location'
            hdr_cells[2].text = 'Characters'
            hdr_cells[3].text = 'Objective'
            hdr_cells[4].text = 'Duration'
            
            for s in scenes:
                row_cells = table.add_row().cells
                row_cells[0].text = str(s.get("scene_number", ""))
                row_cells[1].text = s.get("location", "")
                
                chars_val = s.get("characters", "")
                if isinstance(chars_val, list):
                    chars_str = ", ".join([str(c) for c in chars_val])
                else:
                    chars_str = str(chars_val or "")
                    
                row_cells[2].text = chars_str
                row_cells[3].text = s.get("objective", "")
                row_cells[4].text = s.get("duration", "")
        else:
            doc.add_paragraph("No scenes broken down.")

        # 5. Storyboard
        doc.add_heading("5. Storyboard Frames", level=1)
        sbs = data.get("storyboard", {}).get("storyboards", [])
        for sb in sbs:
            doc.add_heading(f"Scene {sb.get('scene_number', 'N/A')}", level=2)
            doc.add_paragraph(f"Camera Angle: {sb.get('camera_angle', 'N/A')}")
            doc.add_paragraph(f"Lighting: {sb.get('lighting', 'N/A')}")
            doc.add_paragraph(f"Mood: {sb.get('mood', 'N/A')}")
            doc.add_paragraph(f"Prompt: {sb.get('prompt', 'N/A')}")

        # 6. Sound Design
        doc.add_heading("6. Sound Design", level=1)
        sound = data.get("sound_design", {})
        doc.add_paragraph(f"Background Music: {sound.get('background_music', 'N/A')}")
        doc.add_paragraph(f"Ambience: {sound.get('ambience', 'N/A')}")
        doc.add_paragraph(f"Foley: {sound.get('foley_effects', 'N/A')}")
        doc.add_paragraph(f"Dialogue Treatment: {sound.get('dialogue_treatment', 'N/A')}")
        doc.add_paragraph(f"Notes: {sound.get('scene_sound_notes', 'N/A')}")

        # 7. Production Plan
        doc.add_heading("7. Production Plan", level=1)
        prod = data.get("production_plan", {})
        doc.add_paragraph(f"Locations: {prod.get('shooting_locations', 'N/A')}")
        doc.add_paragraph(f"Props: {prod.get('required_props', 'N/A')}")
        doc.add_paragraph(f"Equipment: {prod.get('equipment', 'N/A')}")
        doc.add_paragraph(f"Crew: {prod.get('crew_suggestions', 'N/A')}")
        doc.add_paragraph(f"Est Shoot Days: {prod.get('estimated_shoot_days', 'N/A')}")

        # 8. Budget Planning
        doc.add_heading("8. Budget Planning", level=1)
        budget = data.get("budget_plan", {})
        if budget:
            pre = budget.get("pre_production", {}) or {}
            prod_b = budget.get("production", {}) or {}
            post = budget.get("post_production", {}) or {}
            
            doc.add_heading("Pre-Production", level=2)
            doc.add_paragraph(f"Cost: {pre.get('cost', 'N/A')}")
            doc.add_paragraph(pre.get('details', 'N/A'))
            
            doc.add_heading("Production", level=2)
            doc.add_paragraph(f"Cost: {prod_b.get('cost', 'N/A')}")
            doc.add_paragraph(prod_b.get('details', 'N/A'))
            
            doc.add_heading("Post-Production", level=2)
            doc.add_paragraph(f"Cost: {post.get('cost', 'N/A')}")
            doc.add_paragraph(post.get('details', 'N/A'))
            
            doc.add_paragraph(f"Total Budget: {budget.get('total_budget', 'N/A')}")
            doc.add_paragraph(f"Cost Saving Tips: {budget.get('cost_saving_tips', 'N/A')}")
        else:
            doc.add_paragraph("No budget details generated.")

        # 9. Screenplay
        doc.add_heading("9. Screenplay", level=1)
        screenplay = safe_str(data.get("screenplay", {}).get("screenplay_text", ""))
        if screenplay:
            lines = screenplay.split("\n")
            for line in lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                stripped = line.strip()
                if stripped.startswith("INT.") or stripped.startswith("EXT."):
                    run = p.add_run(stripped)
                    run.bold = True
                elif stripped.isupper() and len(stripped) < 30:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run(stripped)
                elif stripped.startswith("("):
                    p.paragraph_format.left_indent = Inches(1.5)
                    p.add_run(stripped)
                elif len(line) - len(line.lstrip()) > 10:
                    p.paragraph_format.left_indent = Inches(1.0)
                    p.add_run(stripped)
                else:
                    p.add_run(stripped)
        else:
            doc.add_paragraph("No screenplay script compiled.")

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_txt(self, project, data):
        """
        Generates a readable Plain Text (TXT) document.
        Returns bytes.
        """
        out = []
        out.append("="*60)
        out.append(f"CINEFORGE AI PRE-PRODUCTION STUDO BLUEPRINT")
        out.append(f"Project Name: {project.get('project_name', 'Untitled')}")
        out.append(f"Genre: {project.get('genre', 'N/A')}")
        out.append(f"Target Audience: {project.get('target_audience', 'N/A')}")
        out.append(f"Generated On: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
        out.append("="*60 + "\n\n")

        # Story Analysis
        out.append("1. STORY ANALYSIS\n" + "-"*17)
        analysis = data.get("story_analysis", {})
        out.append(f"Logline: {safe_str(analysis.get('logline', 'N/A'))}")
        out.append(f"Tagline: {safe_str(analysis.get('tagline', 'N/A'))}")
        out.append(f"\nSynopsis:\n{safe_str(analysis.get('synopsis', 'N/A'))}")
        out.append(f"\nGenre Analysis:\n{safe_str(analysis.get('genre_analysis', 'N/A'))}")
        out.append(f"\nTheme:\n{safe_str(analysis.get('theme', 'N/A'))}")
        out.append(f"\nAudience Insights:\n{safe_str(analysis.get('audience_insights', 'N/A'))}\n\n")

        # Narrative Structure
        out.append("2. NARRATIVE STRUCTURE\n" + "-"*22)
        struct = data.get("narrative_structure", {})
        for act_key in ["act_1", "act_2", "act_3"]:
            act = struct.get(act_key, {})
            if act:
                out.append(f"\n{safe_str(act.get('title', act_key.replace('_', ' ').title()))}:")
                out.append(f"  Description: {safe_str(act.get('description', 'N/A'))}")
                out.append(f"  Conflict: {safe_str(act.get('conflict', 'N/A'))}")
                out.append(f"  Rising Action: {safe_str(act.get('rising_action', 'N/A'))}")
                if "resolution" in act:
                    out.append(f"  Resolution: {safe_str(act.get('resolution', 'N/A'))}")
        out.append("\n\n")

        # Characters
        out.append("3. CHARACTER PROFILES\n" + "-"*21)
        chars = data.get("characters", {}).get("characters", [])
        for char in chars:
            out.append(f"\n* Name: {safe_str(char.get('name', 'N/A'))} (Age: {safe_str(char.get('age', 'N/A'))})")
            out.append(f"  Backstory: {safe_str(char.get('backstory', 'N/A'))}")
            out.append(f"  Personality: {safe_str(char.get('personality', 'N/A'))}")
            out.append(f"  Goals: {safe_str(char.get('goals', 'N/A'))}")
            out.append(f"  Strengths/Weaknesses: {safe_str(char.get('strengths', 'N/A'))} / {safe_str(char.get('weaknesses', 'N/A'))}")
            out.append(f"  Arc: {safe_str(char.get('arc', 'N/A'))}")
        out.append("\n\n")

        # Scenes
        out.append("4. SCENE BREAKDOWN\n" + "-"*18)
        scenes = data.get("scene_breakdown", {}).get("scenes", [])
        for s in scenes:
            out.append(f"\nScene {s.get('scene_number', 'N/A')}: {s.get('location', 'N/A')}")
            chars_val = s.get("characters", "")
            if isinstance(chars_val, list):
                chars_str = ", ".join([str(c) for c in chars_val])
            else:
                chars_str = str(chars_val or "N/A")
            out.append(f"  Characters: {chars_str}")
            out.append(f"  Objective: {s.get('objective', 'N/A')}")
            out.append(f"  Duration: {s.get('duration', 'N/A')}")
        out.append("\n\n")

        # Storyboard
        out.append("5. STORYBOARD FRAME PLAN\n" + "-"*25)
        sbs = data.get("storyboard", {}).get("storyboards", [])
        for sb in sbs:
            out.append(f"\nScene {sb.get('scene_number', 'N/A')}:")
            out.append(f"  Angle: {sb.get('camera_angle', 'N/A')}")
            out.append(f"  Lighting: {sb.get('lighting', 'N/A')}")
            out.append(f"  Mood: {sb.get('mood', 'N/A')}")
            out.append(f"  Prompt: {sb.get('prompt', 'N/A')}")
        out.append("\n\n")

        # Sound
        out.append("6. SOUND DESIGN\n" + "-"*15)
        sound = data.get("sound_design", {})
        out.append(f"Background Music: {sound.get('background_music', 'N/A')}")
        out.append(f"Ambience: {sound.get('ambience', 'N/A')}")
        out.append(f"Foley Effects: {sound.get('foley_effects', 'N/A')}")
        out.append(f"Dialogue Treatment: {sound.get('dialogue_treatment', 'N/A')}")
        out.append(f"Sound Cue Notes: {sound.get('scene_sound_notes', 'N/A')}\n\n")

        # Production Plan
        out.append("7. PRODUCTION PLANNING\n" + "-"*22)
        prod = data.get("production_plan", {})
        out.append(f"Shooting Locations: {prod.get('shooting_locations', 'N/A')}")
        out.append(f"Required Props: {prod.get('required_props', 'N/A')}")
        out.append(f"Equipment Suggested: {prod.get('equipment', 'N/A')}")
        out.append(f"Crew: {prod.get('crew_suggestions', 'N/A')}")
        out.append(f"Est Shoot Days: {prod.get('estimated_shoot_days', 'N/A')}\n\n")

        # Budget Planning
        out.append("8. BUDGET PLANNING\n" + "-"*18)
        budget = data.get("budget_plan", {})
        if budget:
            pre = budget.get("pre_production", {}) or {}
            prod_b = budget.get("production", {}) or {}
            post = budget.get("post_production", {}) or {}
            out.append(f"Pre-Production Cost: {pre.get('cost', 'N/A')}\nDetails: {pre.get('details', 'N/A')}\n")
            out.append(f"Production Cost: {prod_b.get('cost', 'N/A')}\nDetails: {prod_b.get('details', 'N/A')}\n")
            out.append(f"Post-Production Cost: {post.get('cost', 'N/A')}\nDetails: {post.get('details', 'N/A')}\n")
            out.append(f"Total Budget: {budget.get('total_budget', 'N/A')}\n")
            out.append(f"Saving Tips: {budget.get('cost_saving_tips', 'N/A')}\n\n")
        else:
            out.append("No budget details generated.\n\n")

        # Screenplay
        out.append("9. SCREENPLAY\n" + "-"*13)
        screenplay = safe_str(data.get("screenplay", {}).get("screenplay_text", ""))
        out.append(screenplay)
        
        return "\n".join(out).encode('utf-8')

# Instantiate service singleton
export_service = ExportService()
